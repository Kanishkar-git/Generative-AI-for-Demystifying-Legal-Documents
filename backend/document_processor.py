import io
import re
from pathlib import Path
from typing import Optional
import logging

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes."""
    try:
        from pypdf import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        texts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                texts.append(text)
        return "\n\n".join(texts)
    except Exception as e:
        logger.error(f"PDF extraction error: {e}")
        raise ValueError(f"Could not extract text from PDF: {str(e)}")


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        texts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                texts.append(paragraph.text)
        # Also extract table text
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    texts.append(row_text)
        return "\n\n".join(texts)
    except Exception as e:
        logger.error(f"DOCX extraction error: {e}")
        raise ValueError(f"Could not extract text from DOCX: {str(e)}")


def extract_text_from_txt(file_bytes: bytes) -> str:
    """Extract text from plain text bytes."""
    try:
        return file_bytes.decode("utf-8", errors="replace")
    except Exception as e:
        raise ValueError(f"Could not read text file: {str(e)}")


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Route extraction based on file extension."""
    ext = Path(filename).suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    elif ext == ".txt":
        return extract_text_from_txt(file_bytes)
    else:
        raise ValueError(f"Unsupported file type: {ext}. Supported: PDF, DOCX, TXT")


def clean_text(text: str) -> str:
    """Clean and normalize extracted text."""
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r' {2,}', ' ', text)
    # Remove non-printable characters
    text = re.sub(r'[^\x20-\x7E\n\r\t]', ' ', text)
    return text.strip()


def chunk_text(text: str, chunk_size: int = 2000, overlap: int = 200) -> list[str]:
    """Split text into overlapping chunks for RAG."""
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    while start < len(text):
        end = start + chunk_size
        if end < len(text):
            # Try to break at a sentence boundary
            boundary = text.rfind('. ', start, end)
            if boundary > start + chunk_size // 2:
                end = boundary + 1
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        start = end - overlap
        if start >= len(text):
            break
    return chunks


def process_document(file_bytes: bytes, filename: str) -> dict:
    """Full pipeline: extract → clean → chunk."""
    raw_text = extract_text(file_bytes, filename)
    clean = clean_text(raw_text)
    
    if not clean or len(clean.strip()) < 10:
        raise ValueError("Could not extract any meaningful text from this document. It might be scanned or encrypted.")
        
    chunks = chunk_text(clean)
    
    return {
        "filename": filename,
        "full_text": clean,
        "chunks": chunks,
        "char_count": len(clean),
        "word_count": len(clean.split()),
        "chunk_count": len(chunks),
    }
